import pathlib
import re
from docx import Document

def load_docx(path: str | pathlib.Path) -> tuple[str, dict]:
    """
    Extracts text from a .docx file with structural markers between
    paragraphs and tables for better downstream chunking.

    Returns:
        (text, metadata)
    """
    if not isinstance(path, pathlib.Path):
        path = pathlib.Path(path)

    document = Document(path)
    text_parts = []

    # Process paragraphs
    for paragraph in document.paragraphs:
        clean_para = paragraph.text.strip()
        if clean_para:
            text_parts.append(clean_para)
            text_parts.append("")  # <- for \n\n break after paragraph

    # Process tables
    for table in document.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_rows.append(" | ".join(row_cells))  # Delimit cells in row

        if table_rows:
            text_parts.append("--- TABLE START ---")  # Optional structural marker
            text_parts.extend(table_rows)
            text_parts.append("--- TABLE END ---")
            text_parts.append("")  # <- extra break after table

    # Join with double newlines to support paragraph-based chunking
    full_text = "\n\n".join(text_parts)

    metadata = {"source": str(path), "content_type": "docx", "doc_type": "docx"}
    return full_text, metadata
